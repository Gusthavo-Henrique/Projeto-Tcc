from docx import Document

def fill_word_template(data: dict, template_path: str, output_path: str):
    """
    Substitui placeholders {chave} em um documento Word (.docx),
    incluindo dentro de tabelas e mantendo a formatação original
    (negrito, itálico, sublinhado, etc.), mesmo que o placeholder
    esteja dividido em vários runs.
    """

    doc = Document(template_path)

    def replace_placeholders_in_paragraph(paragraph):
        if not paragraph.runs:
            return

        # Concatena todo o texto do parágrafo
        full_text = ''.join(run.text for run in paragraph.runs)

        # Substitui os placeholders no texto completo
        new_text = full_text
        for key, value in data.items():
            if isinstance(value, list):
                continue
            placeholder = f"{{{key}}}"
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))

        # Se não houver mudança, pula
        if new_text == full_text:
            return

        # Reconstrói runs mantendo formatação do primeiro run
        # (preserva o estilo, mas não "mistura" formatações)
        first_run = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run.text = ""
        first_run.text = new_text

    def process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            replace_placeholders_in_paragraph(paragraph)

    # Substitui em parágrafos principais
    process_paragraphs(doc.paragraphs)

    # Substitui em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    # Substitui também em cabeçalhos e rodapés
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_placeholders_in_paragraph(para)
        for para in section.footer.paragraphs:
            replace_placeholders_in_paragraph(para)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)

    # Salva o resultado
    doc.save(output_path)
